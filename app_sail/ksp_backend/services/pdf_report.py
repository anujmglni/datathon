import os
import time
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

REPORTS_DIR = Path(__file__).resolve().parent.parent / "static" / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def generate_pdf_report(title: str, markdown_content: str, filename: str = None) -> dict:
    """
    Renders styled official executive PDF report using ReportLab.
    """
    timestamp = int(time.time())
    output_filename = filename or f"ksp_report_{timestamp}.pdf"
    filepath = REPORTS_DIR / output_filename

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors

        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'KSPTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=6
        )
        subtitle_style = ParagraphStyle(
            'KSPSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#dc2626'),
            spaceAfter=8
        )
        body_style = ParagraphStyle(
            'KSPBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=6
        )
        bullet_style = ParagraphStyle(
            'KSPBullet',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#1e293b'),
            leftIndent=12,
            spaceAfter=4
        )
        heading_style = ParagraphStyle(
            'KSPHeading',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor('#1d4ed8'),
            spaceBefore=10,
            spaceAfter=6
        )
        subheading_style = ParagraphStyle(
            'KSPSubHeading',
            parent=styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=14,
            textColor=colors.HexColor('#334155'),
            spaceBefore=8,
            spaceAfter=4
        )

        import re

        def convert_md_to_html(text: str) -> str:
            # Escape XML entities
            s = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # Convert **bold** to <b>bold</b>
            s = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', s)
            # Convert `code` to <font fontName="Courier">code</font>
            s = re.sub(r'`(.*?)`', r'<font fontName="Courier">\1</font>', s)
            return s

        story = []

        # Parse markdown lines
        for line in markdown_content.split("\n"):
            line_str = line.strip()
            if not line_str:
                story.append(Spacer(1, 4))
                continue

            if line_str == "---" or line_str == "***":
                story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=6, spaceAfter=8))
            elif line_str.startswith("# "):
                clean_text = convert_md_to_html(line_str[2:].strip())
                story.append(Paragraph(clean_text, title_style))
            elif line_str.startswith("## "):
                clean_text = convert_md_to_html(line_str[3:].strip())
                story.append(Paragraph(clean_text, subtitle_style))
            elif line_str.startswith("### "):
                clean_text = convert_md_to_html(line_str[4:].strip())
                story.append(Paragraph(clean_text, heading_style))
            elif line_str.startswith("#### "):
                clean_text = convert_md_to_html(line_str[5:].strip())
                story.append(Paragraph(clean_text, subheading_style))
            elif line_str.startswith("![") and "](" in line_str:
                try:
                    import base64
                    import io
                    import re
                    from reportlab.platypus import Image as RLImage
                    img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line_str)
                    if img_match:
                        caption = img_match.group(1)
                        img_src = img_match.group(2)
                        if img_src.startswith("data:image"):
                            header, encoded = img_src.split(",", 1)
                            img_data = base64.b64decode(encoded)
                            img_stream = io.BytesIO(img_data)
                            try:
                                from PIL import Image as PILImage
                                pil_img = PILImage.open(img_stream)
                                orig_w, orig_h = pil_img.size
                                aspect = float(orig_h) / float(orig_w) if orig_w > 0 else 0.5
                                target_w = 460
                                target_h = int(target_w * aspect)
                                img_stream.seek(0)
                                rl_img = RLImage(img_stream, width=target_w, height=target_h)
                            except Exception:
                                rl_img = RLImage(img_stream, width=440, height=220)

                            story.append(rl_img)
                            if caption:
                                story.append(Paragraph(f"<i>Figure: {caption}</i>", subtitle_style))
                            story.append(Spacer(1, 8))
                except Exception as img_err:
                    logger.debug(f"PDF Image insert error: {img_err}")
            elif line_str.startswith("- ") or line_str.startswith("* "):
                clean_text = convert_md_to_html(line_str[2:].strip())
                story.append(Paragraph(f"• {clean_text}", bullet_style))
            else:
                clean_text = convert_md_to_html(line_str)
                story.append(Paragraph(clean_text, body_style))

        doc.build(story)
        file_size = filepath.stat().st_size if filepath.exists() else 0

        return {
            "success": True,
            "filename": output_filename,
            "filepath": str(filepath),
            "size_bytes": file_size,
            "download_url": f"/static/reports/{output_filename}"
        }

    except Exception as e:
        logger.debug(f"ReportLab PDF generation fallback error: {e}")
        return {
            "success": False,
            "error": str(e),
            "filename": output_filename
        }


def generate_docx_report(title: str, markdown_content: str, filename: str = None) -> dict:
    """
    Renders styled official executive DOCX report using python-docx.
    """
    timestamp = int(time.time())
    output_filename = filename or f"ksp_report_{timestamp}.docx"
    filepath = REPORTS_DIR / output_filename

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Document Header
        hdr_p = doc.add_paragraph()
        hdr_run = hdr_p.add_run("KARNATAKA STATE POLICE — CRIME INTELLIGENCE PLATFORM")
        hdr_run.font.name = "Arial"
        hdr_run.font.size = Pt(9)
        hdr_run.font.bold = True
        hdr_run.font.color.rgb = RGBColor(71, 85, 105)

        # Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title or "Executive Investigation Briefing")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(30, 58, 138)
        title_p.paragraph_format.space_after = Pt(12)

        # Content parsing
        for line in markdown_content.split("\n"):
            line_str = line.strip()
            if not line_str:
                continue

            if line_str.startswith("#"):
                clean_heading = line_str.lstrip("#").strip().replace("**", "").replace("`", "")
                h_p = doc.add_paragraph()
                h_run = h_p.add_run(clean_heading)
                h_run.font.name = "Arial"
                h_run.font.size = Pt(13)
                h_run.font.bold = True
                h_run.font.color.rgb = RGBColor(29, 78, 216)
                h_p.paragraph_format.space_before = Pt(10)
                h_p.paragraph_format.space_after = Pt(4)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                clean_bullet = line_str[2:].strip().replace("**", "").replace("`", "")
                b_p = doc.add_paragraph(style='List Bullet')
                b_run = b_p.add_run(clean_bullet)
                b_run.font.name = "Arial"
                b_run.font.size = Pt(10)
                b_run.font.color.rgb = RGBColor(15, 23, 42)
            else:
                clean_body = line_str.replace("**", "").replace("`", "")
                p = doc.add_paragraph()
                p_run = p.add_run(clean_body)
                p_run.font.name = "Arial"
                p_run.font.size = Pt(10)
                p_run.font.color.rgb = RGBColor(15, 23, 42)
                p.paragraph_format.space_after = Pt(6)

        doc.save(str(filepath))
        file_size = filepath.stat().st_size if filepath.exists() else 0

        return {
            "success": True,
            "filename": output_filename,
            "filepath": str(filepath),
            "size_bytes": file_size,
            "download_url": f"/static/reports/{output_filename}"
        }
    except Exception as e:
        logger.debug(f"python-docx report generation error: {e}")
        return {
            "success": False,
            "error": str(e),
            "filename": output_filename
        }


